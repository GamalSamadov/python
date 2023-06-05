"""
The add_paragraph() function is used to add a paragraph to the file.

Text alignment can be specified through the alignment property:
https://python-docx.readthedocs.io/en/latest/api/enum/WdAlignParagraph.html#wdparagraphalignment

The style property is used to modify the formatting of a paragraph.
The delete() function is used to delete a paragraph formatting.
The add_picture() function adds an image to the file.
"""

import docx
from pathlib import Path
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

myDoc = docx.Document()
# myDoc.add_paragraph('This is a pharagraph written on Python')
# myDoc.add_paragraph('This is a second pharagraph written on Python')
#
# # إضافة النص باللغة العربية
# paragraph = myDoc.add_paragraph('نص عربي من البايثون')
# paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
#
# thirdParagraph = myDoc.add_paragraph('This is a third paragraph. ')
# thirdParagraph.add_run('Second section of the third paragraph.')
#
# # Add heading
# myDoc.add_heading('This is a level 1 Heading', 0)
# myDoc.add_heading('This is a level 2 Heading', 1)
# myDoc.add_heading('This is a level 3 Heading', 2)
# myDoc.add_heading('This is a level 4 Heading', 3)
# myDoc.add_heading('This is a level 5 Heading', 4)
#
# # Style
# print(myDoc.paragraphs[0].style)
# print(myDoc.paragraphs[5].style)
# print(myDoc.paragraphs[4].style)
#
# myDoc.paragraphs[0].style = myDoc.styles['Heading 1']
# myDoc.paragraphs[3].style = myDoc.styles['Heading 4']
# myDoc.paragraphs[3].style.delete()
#
# myDoc.add_paragraph('This is a styled paragraph', 'Title')

# Add Pictures
# pth = str(Path.home() / Path('OneDrive', 'Desktop', 'Curces', 'Hsoub Academy', 'The Code', '02 Practical applications using Python', '08_2 Word', 'word_files', 'test.png'))
# myDoc.add_picture(pth, width=docx.shared.Inches(5), height=docx.shared.Inches(7))
#
# pth = Path.home() / Path('OneDrive', 'Desktop', 'Curces', 'Hsoub Academy', 'The Code', '02 Practical applications using Python', '08_2 Word')
# myDoc.save(pth / Path('write.docx'))


# Add Pictures
# pth = str(Path.home() / Path('Desktop', "france.jpg"))
pth = r"C:\Users\gamal\OneDrive\Desktop\france.jpg"

print(pth)

myDoc.add_picture(pth, width=docx.shared.Inches(5), height=docx.shared.Inches(7))

pth = str(Path.home() / Path('Desktop'))
print(pth)
# done = myDoc.save(Path.home() / Path('Desktop', 'write.docx'))
done = myDoc.save(r'C:\Users\gamal\OneDrive\Desktop\write1.docx')

print(done)