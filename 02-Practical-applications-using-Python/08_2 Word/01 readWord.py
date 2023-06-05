"""
Link to the official site to install Microsoft Office.
docx package install:
pip install python-docx
The Document() builder returns an object representing the specified Word file and allows you to manipulate and manipulate that file.
The paragraphs property returns a list of objects that represent the paragraphs within a Word file.
"""

import docx
from pathlib import Path
from readDocx import get_text

pth = Path.home() / Path('OneDrive', 'Desktop', 'Curces', 'Hsoub Academy', 'The Code', '02 Practical applications using Python', '08_2 Word')
doc = docx.Document(pth / Path('word_files', 'academy_1.docx'))

print(len(doc.paragraphs))
print(doc.paragraphs[0].text)
print(doc.paragraphs[1].text)
print(doc.paragraphs[2].text)

print(len(doc.paragraphs[2].runs))
print(doc.paragraphs[2].runs[0].text)
print(doc.paragraphs[2].runs[1].text)
print(doc.paragraphs[2].runs[2].text)
print(doc.paragraphs[2].runs[3].text)
print(doc.paragraphs[2].runs[4].text)

print('-'*50)

pth1 = pth / Path('word_files', 'academy_1.docx')
print(get_text(pth1))